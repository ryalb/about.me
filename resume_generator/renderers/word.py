"""Word (.docx) renderer using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Pt, RGBColor

from ..contact import profile_display
from ..i18n import present_label, section_labels, work_cutoff_notice

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

# ── Colour palette ────────────────────────────────────────────────────────
_ACCENT = RGBColor(0x25, 0x63, 0xEB)  # #2563EB blue
_MUTED = RGBColor(0x64, 0x74, 0x8B)  # #64748B slate
_TEXT = RGBColor(0x1E, 0x29, 0x3B)  # #1E293B dark


def _pt(size: float, zoom: float) -> Length:
    """Scale a typographic length (in points) by the content zoom factor.

    Applies to font sizes and paragraph spacing.  Page margins and indents are
    page geometry and stay fixed, so zoom changes how much text fits a page
    rather than the shape of the page — matching the html/pdf renderers.
    """
    return Pt(size * zoom)


def _set_font(
    run,
    bold: bool = False,
    italic: bool = False,
    size: int | None = None,
    color: RGBColor | None = None,
    *,
    zoom: float = 1.0,
) -> None:
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = _pt(size, zoom)
    if color:
        run.font.color.rgb = color


def _add_heading(
    doc: DocxDocument, text: str, level: int = 1, *, zoom: float = 1.0
) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = _pt(12 if level == 1 else 8, zoom)
    para.paragraph_format.space_after = _pt(4 if level == 1 else 2, zoom)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = _ACCENT if level > 1 else _TEXT
    run.font.size = _pt({1: 20, 2: 13, 3: 11}.get(level, 11), zoom)
    if level == 2:
        # Add a bottom border to simulate a section divider
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CBD5E1")
        pBdr.append(bottom)
        pPr.append(pBdr)


def _add_hyperlink(
    para,
    url: str,
    text: str,
    size: float = 10.5,
    bold: bool = True,
    *,
    zoom: float = 1.0,
) -> None:
    """Append *text* to *para* as a real Word hyperlink pointing at *url*.

    python-docx has no public hyperlink API, so the relationship and the
    ``w:hyperlink`` element are created by hand.
    """
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(round(size * zoom * 2)))  # half-points
    rPr.append(sz)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    para._p.append(link)


def _add_entry_header(
    doc: DocxDocument,
    title: str,
    subtitle: str = "",
    date_str: str = "",
    url: str = "",
    *,
    zoom: float = 1.0,
) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = _pt(6, zoom)
    para.paragraph_format.space_after = Pt(0)

    # Title — hyperlinked when the entry carries a url
    if url:
        _add_hyperlink(para, url, title, zoom=zoom)
    else:
        r = para.add_run(title)
        r.bold = True
        r.font.size = _pt(10.5, zoom)
        r.font.color.rgb = _TEXT

    if subtitle:
        r2 = para.add_run(f"  {subtitle}")
        r2.font.size = _pt(10, zoom)
        r2.font.color.rgb = _MUTED

    # Right-aligned date via tab stop
    if date_str:
        tab = para.add_run("\t" + date_str)
        tab.font.size = _pt(9, zoom)
        tab.font.color.rgb = _MUTED
        tab.italic = True
        # Set a right-aligned tab stop at page width
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab_el = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:pos"), "9360")  # ~6.5 inches in twips
        tabs.append(tab_el)
        pPr.append(tabs)


def _add_body_text(doc: DocxDocument, text: str, *, zoom: float = 1.0) -> None:
    if not text:
        return
    para = doc.add_paragraph(text)
    para.paragraph_format.space_before = _pt(2, zoom)
    para.paragraph_format.space_after = _pt(2, zoom)
    para.paragraph_format.left_indent = Inches(0.1)
    for run in para.runs:
        run.font.size = _pt(9.5, zoom)
        run.font.color.rgb = _TEXT


def _add_bullet(doc: DocxDocument, text: str, *, zoom: float = 1.0) -> None:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.size = _pt(9.5, zoom)
    run.font.color.rgb = _TEXT
    para.paragraph_format.space_before = _pt(1, zoom)
    para.paragraph_format.space_after = _pt(1, zoom)
    para.paragraph_format.left_indent = Inches(0.2)


def _date_range(start: str | None, end: str | None, present: str) -> str:
    if not start and not end:
        return ""
    if start and end:
        return f"{start} – {end}"
    if start:
        return f"{start} – {present}"
    return end or ""


def render_word(resume: dict[str, Any], output_path: Path, zoom: float = 1.0) -> None:
    """Render a JSON Resume dict to a Word .docx file.

    *zoom* scales font sizes and paragraph spacing (1.0 = unscaled).
    """
    doc = Document()
    present = present_label(resume)
    labels = section_labels(resume)

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Basics ────────────────────────────────────────────────────────────
    basics = resume.get("basics") or {}
    name = basics.get("name") or "Resume"

    _add_heading(doc, name, level=1, zoom=zoom)

    if label := basics.get("label"):
        p = doc.add_paragraph(label)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = _pt(2, zoom)
        for r in p.runs:
            r.font.size = _pt(12, zoom)
            r.font.color.rgb = _ACCENT

    # Contact line
    contact_parts: list[str] = []
    if email := basics.get("email"):
        contact_parts.append(email)
    if phone := basics.get("phone"):
        contact_parts.append(phone)
    if url := basics.get("url"):
        contact_parts.append(url)
    if loc := basics.get("location"):
        parts = [loc.get("city"), loc.get("region"), loc.get("countryCode")]
        loc_str = ", ".join(p for p in parts if p)
        if loc_str:
            contact_parts.append(loc_str)
    for profile in basics.get("profiles") or []:
        if shown := profile_display(profile):
            contact_parts.append(shown)

    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = _pt(4, zoom)
        for r in p.runs:
            r.font.size = _pt(9, zoom)
            r.font.color.rgb = _MUTED

    if summary := basics.get("summary"):
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = _pt(6, zoom)
        for r in p.runs:
            r.font.size = _pt(10, zoom)
            r.font.color.rgb = _TEXT

    # ── Skills ────────────────────────────────────────────────────────────
    if skills := resume.get("skills"):
        _add_heading(doc, labels["skills"], level=2, zoom=zoom)
        for skill in skills:
            sname = skill.get("name") or ""
            level = skill.get("level") or ""
            kws = skill.get("keywords") or []
            p = doc.add_paragraph()
            p.paragraph_format.space_before = _pt(2, zoom)
            p.paragraph_format.space_after = _pt(2, zoom)
            r = p.add_run(sname)
            r.bold = True
            r.font.size = _pt(10, zoom)
            if level:
                r2 = p.add_run(f" ({level})")
                r2.font.size = _pt(9.5, zoom)
                r2.font.color.rgb = _MUTED
            if kws:
                r3 = p.add_run(": " + ", ".join(kws))
                r3.font.size = _pt(9.5, zoom)
                r3.font.color.rgb = _TEXT

    # ── Work ──────────────────────────────────────────────────────────────
    if work := resume.get("work"):
        _add_heading(doc, labels["work"], level=2, zoom=zoom)
        for job in work:
            pos = job.get("position") or ""
            company = job.get("name") or ""
            loc = job.get("location") or ""
            subtitle = " | ".join(filter(None, [company, loc]))
            dr = _date_range(job.get("startDate"), job.get("endDate"), present)
            _add_entry_header(doc, pos, subtitle=subtitle, date_str=dr, zoom=zoom)
            if s := job.get("summary"):
                _add_body_text(doc, s, zoom=zoom)
            for h in job.get("highlights") or []:
                _add_bullet(doc, h, zoom=zoom)
        if notice := work_cutoff_notice(resume):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = _pt(6, zoom)
            para.paragraph_format.space_after = _pt(2, zoom)
            run = para.add_run(notice)
            run.italic = True
            run.font.size = _pt(9, zoom)
            run.font.color.rgb = _MUTED

    # ── Projects ──────────────────────────────────────────────────────────
    if projects := resume.get("projects"):
        _add_heading(doc, labels["projects"], level=2, zoom=zoom)
        for proj in projects:
            pname = proj.get("name") or ""
            ptype = proj.get("type") or ""
            dr = _date_range(proj.get("startDate"), proj.get("endDate"), present)
            _add_entry_header(doc, pname, subtitle=ptype, date_str=dr, zoom=zoom)
            if desc := proj.get("description"):
                _add_body_text(doc, desc, zoom=zoom)
            kws = proj.get("keywords") or []
            if kws:
                _add_body_text(doc, "Keywords: " + ", ".join(kws), zoom=zoom)
            for h in proj.get("highlights") or []:
                _add_bullet(doc, h, zoom=zoom)

    # ── Volunteer ─────────────────────────────────────────────────────────
    if volunteer := resume.get("volunteer"):
        _add_heading(doc, labels["volunteer"], level=2, zoom=zoom)
        for v in volunteer:
            pos = v.get("position") or ""
            org = v.get("organization") or ""
            dr = _date_range(v.get("startDate"), v.get("endDate"), present)
            _add_entry_header(doc, pos, subtitle=org, date_str=dr, zoom=zoom)
            if s := v.get("summary"):
                _add_body_text(doc, s, zoom=zoom)
            for h in v.get("highlights") or []:
                _add_bullet(doc, h, zoom=zoom)

    # ── Education ─────────────────────────────────────────────────────────
    if education := resume.get("education"):
        _add_heading(doc, labels["education"], level=2, zoom=zoom)
        for edu in education:
            degree_parts = [edu.get("studyType"), edu.get("area")]
            degree = " in ".join(p for p in degree_parts if p)
            institution = edu.get("institution") or ""
            dr = _date_range(edu.get("startDate"), edu.get("endDate"), present)
            _add_entry_header(
                doc, degree or "Degree", subtitle=institution, date_str=dr, zoom=zoom
            )
            if score := edu.get("score"):
                _add_body_text(doc, f"Score: {score}", zoom=zoom)
            for c in edu.get("courses") or []:
                _add_bullet(doc, c, zoom=zoom)

    # ── Certificates ──────────────────────────────────────────────────────
    if certs := resume.get("certificates"):
        _add_heading(doc, labels["certificates"], level=2, zoom=zoom)
        for c in certs:
            cname = c.get("name") or ""
            issuer = c.get("issuer") or ""
            cdate = c.get("date") or ""
            _add_entry_header(
                doc,
                cname,
                subtitle=issuer,
                date_str=cdate,
                url=c.get("url") or "",
                zoom=zoom,
            )

    # ── Publications ──────────────────────────────────────────────────────
    if pubs := resume.get("publications"):
        _add_heading(doc, labels["publications"], level=2, zoom=zoom)
        for pub in pubs:
            pname = pub.get("name") or ""
            publisher = pub.get("publisher") or ""
            pdate = pub.get("releaseDate") or ""
            _add_entry_header(
                doc,
                pname,
                subtitle=publisher,
                date_str=pdate,
                url=pub.get("url") or "",
                zoom=zoom,
            )
            if s := pub.get("summary"):
                _add_body_text(doc, s, zoom=zoom)

    # ── Awards ────────────────────────────────────────────────────────────
    if awards := resume.get("awards"):
        _add_heading(doc, labels["awards"], level=2, zoom=zoom)
        for a in awards:
            title = a.get("title") or ""
            awarder = a.get("awarder") or ""
            adate = a.get("date") or ""
            dr = adate
            _add_entry_header(
                doc,
                title,
                subtitle=awarder,
                date_str=dr,
                url=a.get("url") or "",
                zoom=zoom,
            )
            if s := a.get("summary"):
                _add_body_text(doc, s, zoom=zoom)

    # ── Languages ─────────────────────────────────────────────────────────
    if langs := resume.get("languages"):
        _add_heading(doc, labels["languages"], level=2, zoom=zoom)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = _pt(2, zoom)
        parts = []
        for lang in langs:
            lname = lang.get("language") or ""
            fluency = lang.get("fluency") or ""
            parts.append(lname + (f" ({fluency})" if fluency else ""))
        r = p.add_run(" | ".join(parts))
        r.font.size = _pt(10, zoom)

    # ── Interests ─────────────────────────────────────────────────────────
    if interests := resume.get("interests"):
        _add_heading(doc, labels["interests"], level=2, zoom=zoom)
        for interest in interests:
            iname = interest.get("name") or ""
            kws = interest.get("keywords") or []
            p = doc.add_paragraph()
            r = p.add_run(iname)
            r.bold = True
            r.font.size = _pt(10, zoom)
            if kws:
                r2 = p.add_run(": " + ", ".join(kws))
                r2.font.size = _pt(9.5, zoom)
                r2.font.color.rgb = _TEXT

    # ── References ────────────────────────────────────────────────────────
    if refs := resume.get("references"):
        _add_heading(doc, labels["references"], level=2, zoom=zoom)
        for ref in refs:
            rname = ref.get("name") or ""
            rtext = ref.get("reference") or ""
            p = doc.add_paragraph()
            r = p.add_run(rname)
            r.bold = True
            r.font.size = _pt(10, zoom)
            if rtext:
                _add_body_text(doc, f'"{rtext}"', zoom=zoom)

    doc.save(str(output_path))
